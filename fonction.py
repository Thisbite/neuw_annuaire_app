from fuzzywuzzy import process
from pymongo import MongoClient
import pandas as pd
from docx import Document

# Connexion à MongoDB (consider adding error handling)
client = MongoClient('mongodb://localhost:27017/')
db = client['word_tables']
collection = db['tables']
collection.create_index([('title', 'text')])


def extract_tables_from_docx(file_path):
    document = Document(file_path)
    tables = []
    for table in document.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    return tables


def search_keyword_in_tables(keyword):
    results = []
    for table in collection.find():
        table_data = table['table_data']
        headers = table_data[0]

        # Rechercher dans les entêtes
        for col_index, header in enumerate(headers):
            if keyword in header:
                column_data = [row[col_index] for row in table_data]
                first_column = [row[0] for row in table_data]
                second_column = [row[1] for row in table_data]
                results.append({'table_id': table['table_id'], 'column_data': column_data, 'first_column': first_column,
                                'second_column': second_column, 'header': header})

        # Rechercher dans les lignes
        for row in table_data:
            if any(keyword in cell for cell in row):
                results.append({'table_id': table['table_id'], 'row': row, 'headers': headers})
    return results



def search_title_with_keyword(keyword):
    results = []

    # Perform a text search on the 'title' field
    matched_titles = collection.find({"$text": {"$search": keyword}})

    # Create a list of titles and table IDs from the text search results
    titles_and_ids = [(title['title'], title['table_id']) for title in matched_titles]

    # Use fuzzywuzzy to find close matches with a score cutoff of 90
    fuzzy_matches = process.extractBests(keyword, [t[0] for t in titles_and_ids], score_cutoff=60)

    # Get corresponding table IDs from the fuzzy matches
    matched_table_ids = [t[1] for t in fuzzy_matches]

    # Retrieve matching tables
    for table_id in matched_table_ids:
        table = collection.find_one({'table_id': table_id})
        if table:
            results.append({'table_id': table_id, 'table_data': table['table_data'], 'title': table['title']})

    return results


def update_titles_from_excel(excel_file_path):
    # Charger le fichier Excel
    df = pd.read_excel(excel_file_path)

    # Mettre à jour les titres dans MongoDB
    for _, row in df.iterrows():
        table_id = row['id']
        title = row['titre']
        collection.update_one({'table_id': table_id}, {'$set': {'title': title}})





def extract_tables_from_docx(file_path):
    document = Document(file_path)
    tables = []
    for table in document.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    return tables

def load_titles_from_excel(excel_file_path):
    df = pd.read_excel(excel_file_path)
    titles = df['titre'].tolist()
    regions = df['region'].tolist()
    return titles, regions

def insert_tables_into_mongodb(tables, titles, regions):
    for idx, table in enumerate(tables):
        title = titles[idx] if idx < len(titles) else f"Table {idx + 1}"
        region = regions[idx] if idx < len(regions) else "Unknown"
        collection.insert_one({'table_id': idx, 'table_data': table, 'title': title, 'region': region})

def execution(docx_file_path, excel_file_path):
    tables = extract_tables_from_docx(docx_file_path)
    titles, regions = load_titles_from_excel(excel_file_path)
    insert_tables_into_mongodb(tables, titles, regions)


def delete_all_records():
    collection.delete_many({})
# Exemple d'utilisation

#delete_all_records()
#docx_file_path = 'AnnuaireTchologo Version Finale.docx'
#excel_file_path = 'excel_titre'
#execution(docx_file_path, excel_file_path)



